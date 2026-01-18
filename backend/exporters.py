"""Export transcript to PDF and DOCX formats"""
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML, CSS


def export_to_docx(
    title: str,
    chapters: list,
    takeaways: list,
    transcript: str,
    font_name: str = "Arial",
    font_size: int = 11,
    thumbnail_path: str = None,
    channel: str = ""
) -> bytes:
    """
    Export transcript to DOCX format with thumbnail.

    Returns:
        bytes of the DOCX file
    """
    doc = Document()

    # Add thumbnail at the top if available
    if thumbnail_path and Path(thumbnail_path).exists():
        try:
            doc.add_picture(thumbnail_path, width=Inches(4))
            # Center the image
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass  # Skip if image can't be added

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Channel/Author
    if channel:
        channel_para = doc.add_paragraph()
        channel_run = channel_para.add_run(f"By: {channel}")
        channel_run.font.size = Pt(12)
        channel_run.font.italic = True
        channel_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Chapters section
    if chapters:
        doc.add_heading("Chapters", level=1)
        for ch in chapters:
            p = doc.add_paragraph()
            run = p.add_run(f"{ch.get('timestamp', '')} - {ch.get('title', '')}")
            run.font.size = Pt(font_size)
            run.font.name = font_name

    # Key Takeaways section
    if takeaways:
        doc.add_heading("Key Takeaways", level=1)
        for i, takeaway in enumerate(takeaways, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {takeaway}")
            run.font.size = Pt(font_size)
            run.font.name = font_name

    # Transcript section
    doc.add_heading("Transcript", level=1)

    # Split transcript into paragraphs
    paragraphs = transcript.split('\n\n') if '\n\n' in transcript else [transcript]
    for para_text in paragraphs:
        if para_text.strip():
            p = doc.add_paragraph()
            run = p.add_run(para_text.strip())
            run.font.size = Pt(font_size)
            run.font.name = font_name

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_to_pdf(
    title: str,
    chapters: list,
    takeaways: list,
    transcript: str,
    font_name: str = "Arial",
    font_size: int = 11,
    highlights: list = None,
    thumbnail_path: str = None,
    channel: str = ""
) -> bytes:
    """
    Export transcript to PDF format using WeasyPrint with thumbnail.

    Args:
        highlights: list of dicts with 'text' and 'color' to highlight
        thumbnail_path: path to thumbnail image
        channel: channel/author name

    Returns:
        bytes of the PDF file
    """
    import base64

    highlights = highlights or []

    # Apply highlights to transcript
    highlighted_transcript = transcript
    for h in highlights:
        text = h.get("text", "")
        color = h.get("color", "#ffff00")
        if text:
            highlighted_transcript = highlighted_transcript.replace(
                text,
                f'<span style="background-color: {color};">{text}</span>'
            )

    # Build thumbnail HTML
    thumbnail_html = ""
    if thumbnail_path and Path(thumbnail_path).exists():
        try:
            with open(thumbnail_path, 'rb') as f:
                img_data = base64.b64encode(f.read()).decode('utf-8')
            thumbnail_html = f'<div class="thumbnail"><img src="data:image/jpeg;base64,{img_data}" alt="Thumbnail"></div>'
        except Exception:
            pass

    # Build channel HTML
    channel_html = f'<p class="channel">By: {channel}</p>' if channel else ""

    # Build HTML
    chapters_html = ""
    if chapters:
        chapters_html = "<h2>Chapters</h2><ul>"
        for ch in chapters:
            chapters_html += f"<li><strong>{ch.get('timestamp', '')}</strong> - {ch.get('title', '')}</li>"
        chapters_html += "</ul>"

    takeaways_html = ""
    if takeaways:
        takeaways_html = "<h2>Key Takeaways</h2><ol>"
        for t in takeaways:
            takeaways_html += f"<li>{t}</li>"
        takeaways_html += "</ol>"

    # Convert newlines to paragraphs
    transcript_paragraphs = highlighted_transcript.replace('\n\n', '</p><p>').replace('\n', '<br>')

    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                size: letter;
                margin: 1in;
            }}
            body {{
                font-family: {font_name}, sans-serif;
                font-size: {font_size}pt;
                line-height: 1.6;
                color: #333;
            }}
            .thumbnail {{
                text-align: center;
                margin-bottom: 20px;
            }}
            .thumbnail img {{
                max-width: 400px;
                border-radius: 8px;
                box-shadow: 0 2px 8px rgba(0,0,0,0.2);
            }}
            h1 {{
                text-align: center;
                color: #1a1a1a;
                border-bottom: 2px solid #0066cc;
                padding-bottom: 10px;
                margin-top: 10px;
            }}
            .channel {{
                text-align: center;
                font-style: italic;
                color: #666;
                margin-bottom: 20px;
            }}
            h2 {{
                color: #0066cc;
                margin-top: 20px;
            }}
            ul, ol {{
                margin-left: 20px;
            }}
            li {{
                margin-bottom: 8px;
            }}
            .transcript {{
                text-align: justify;
                margin-top: 20px;
            }}
            .transcript p {{
                margin-bottom: 12px;
            }}
        </style>
    </head>
    <body>
        {thumbnail_html}
        <h1>{title}</h1>
        {channel_html}
        {chapters_html}
        {takeaways_html}
        <h2>Transcript</h2>
        <div class="transcript">
            <p>{transcript_paragraphs}</p>
        </div>
    </body>
    </html>
    """

    # Generate PDF
    html = HTML(string=html_content)
    pdf_bytes = html.write_pdf()

    return pdf_bytes


if __name__ == "__main__":
    # Test export
    test_data = {
        "title": "Test Video Transcript",
        "chapters": [
            {"timestamp": "00:00", "title": "Introduction"},
            {"timestamp": "05:30", "title": "Main Content"},
        ],
        "takeaways": [
            "First important point",
            "Second important point",
        ],
        "transcript": "This is a test transcript.\n\nIt has multiple paragraphs.\n\nAnd some more content here."
    }

    # Test DOCX
    docx_bytes = export_to_docx(**test_data)
    Path("test_output.docx").write_bytes(docx_bytes)
    print("Created test_output.docx")

    # Test PDF
    pdf_bytes = export_to_pdf(**test_data)
    Path("test_output.pdf").write_bytes(pdf_bytes)
    print("Created test_output.pdf")
